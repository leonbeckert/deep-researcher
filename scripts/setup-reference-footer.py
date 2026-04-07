#!/usr/bin/env python3
"""Add a branded footer to the DOCX reference template.

Adds to the default section footer:
  Left:  "Deep Researcher · deep-research.leon.fm" (8pt, grey)
  Right: page number (8pt, grey)
  Top border: thin grey line as separator

Run once, commit the result:
  python3 scripts/setup-reference-footer.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

TEMPLATE = Path(__file__).resolve().parent.parent / "templates" / "reference.docx"

def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing footer content
    for p in footer.paragraphs:
        p.clear()

    # Remove all existing paragraphs from footer
    footer_elem = footer._element
    for p in footer_elem.findall(qn("w:p")):
        footer_elem.remove(p)

    # Create footer paragraph with top border (separator line)
    p = footer.add_paragraph()
    pPr = p._element.get_or_add_pPr()

    # Top border as separator
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="4" w:color="999999"/>'
        "</w:pBdr>"
    )
    pPr.append(borders)

    # Tab stops: right-aligned at page width (~6.5 inches = 9360 twips)
    tabs = parse_xml(
        f'<w:tabs {nsdecls("w")}>'
        '  <w:tab w:val="right" w:pos="9360"/>'
        "</w:tabs>"
    )
    pPr.append(tabs)

    # Left side: attribution text
    run = p.add_run("Deep Researcher \u00b7 deep-research.leon.fm")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Tab to right side
    tab_run = p.add_run()
    tab_run._element.append(parse_xml(f'<w:tab {nsdecls("w")}/>'))

    # Right side: page number field
    page_run = p.add_run()
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    page_run._element.append(fldChar_begin)

    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    page_run._element.append(instrText)

    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    page_run._element.append(fldChar_end)


if __name__ == "__main__":
    doc = Document(str(TEMPLATE))
    add_footer(doc)
    doc.save(str(TEMPLATE))
    print(f"Footer added to {TEMPLATE}")
